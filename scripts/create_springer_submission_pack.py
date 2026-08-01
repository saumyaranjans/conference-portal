from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "springer_submission_pack"

NAVY = "0B2E63"
BLUE = "1E5AA8"
DARK_BLUE = "153E6F"
GOLD = "C99424"
INK = "1F2937"
MUTED = "5F6B7A"
LIGHT_BLUE = "E8EEF5"
LIGHT_GOLD = "FFF7E3"
LIGHT_GRAY = "F4F6F9"
BORDER = "C9D3E0"
WHITE = "FFFFFF"

GJFSM_GUIDELINES = "https://link.springer.com/journal/40171/submission-guidelines"
GJFSM_SCOPE = "https://link.springer.com/journal/40171/aims-and-scope"
IJGBC_GUIDELINES = "https://link.springer.com/journal/42943/submission-guidelines"
IJGBC_SCOPE = "https://link.springer.com/journal/42943/aims-and-scope"


def set_run_font(run, *, name="Calibri", size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: Sequence[int], indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = tr_pr.find(qn("w:cantSplit"))
        if cant_split is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for index, cell in enumerate(row.cells):
            width = widths_dxa[min(index, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def mark_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def add_hyperlink(paragraph, text: str, url: str, color=BLUE):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def paragraph_border_bottom(paragraph, color=GOLD, size="18", space="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_field(paragraph):
    paragraph.add_run("Page ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for field_node in (begin, instr, separate, value, end):
        field_run = OxmlElement("w:r")
        field_run.append(field_node)
        paragraph._p.append(field_run)


def configure_document(doc: Document, running_label: str):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Title", 25, NAVY, 0, 8),
        ("Subtitle", 12, MUTED, 0, 14),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.0

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    set_run_font(hp.add_run("GLOGIFT 2027  |  "), size=8.5, color=NAVY, bold=True)
    set_run_font(hp.add_run(running_label), size=8.5, color=MUTED)
    paragraph_border_bottom(hp, color=BORDER, size="6", space="3")

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    set_run_font(fp.add_run("GLOGIFT 2027  |  "), size=8, color=MUTED)
    add_page_field(fp)
    for run in fp.runs:
        set_run_font(run, size=8, color=MUTED)


def add_title_block(doc: Document, kicker: str, title: str, subtitle: str, status: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run(kicker.upper()), size=9, color=GOLD, bold=True)
    p = doc.add_paragraph(style="Title")
    p.paragraph_format.space_after = Pt(6)
    set_run_font(p.add_run(title), size=25, color=NAVY, bold=True)
    p = doc.add_paragraph(style="Subtitle")
    p.paragraph_format.space_after = Pt(8)
    set_run_font(p.add_run(subtitle), size=12, color=MUTED)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    set_run_font(p.add_run(status), size=9.5, color=DARK_BLUE, bold=True)
    paragraph_border_bottom(p, color=GOLD, size="18", space="6")


def add_callout(doc: Document, title: str, text: str, *, fill=LIGHT_GOLD, accent=GOLD):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, color=accent, size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_run_font(p.add_run(title + "  "), size=10.5, color=NAVY, bold=True)
    set_run_font(p.add_run(text), size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullet(doc: Document, text: str, *, bold_prefix: str | None = None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(p.add_run(bold_prefix), bold=True)
        set_run_font(p.add_run(text[len(bold_prefix):]))
    else:
        set_run_font(p.add_run(text))
    return p


def add_number(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    set_run_font(p.add_run(text))
    return p


def add_field_table(doc: Document, rows: Iterable[tuple[str, str]], widths=(2700, 6660)):
    data = list(rows)
    table = doc.add_table(rows=len(data), cols=2)
    set_table_geometry(table, list(widths))
    set_table_borders(table)
    for index, (label, value) in enumerate(data):
        left, right = table.rows[index].cells
        set_cell_shading(left, LIGHT_BLUE)
        lp = left.paragraphs[0]
        rp = right.paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        rp.paragraph_format.space_after = Pt(0)
        set_run_font(lp.add_run(label), size=10.5, color=NAVY, bold=True)
        set_run_font(rp.add_run(value), size=10.5, color=INK)
    return table


def add_source_list(doc: Document):
    doc.add_heading("Official journal sources", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run("Global Journal of Flexible Systems Management — Submission guidelines: "), size=9.5, color=MUTED)
    add_hyperlink(p, "Open the official submission guidelines", GJFSM_GUIDELINES)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run("Global Journal of Flexible Systems Management — Aims and scope: "), size=9.5, color=MUTED)
    add_hyperlink(p, "Open the official aims and scope", GJFSM_SCOPE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run("International Journal of Global Business and Competitiveness — Submission guidelines: "), size=9.5, color=MUTED)
    add_hyperlink(p, "Open the official submission guidelines", IJGBC_GUIDELINES)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run("International Journal of Global Business and Competitiveness — Aims and scope: "), size=9.5, color=MUTED)
    add_hyperlink(p, "Open the official aims and scope", IJGBC_SCOPE)


def create_author_details_worksheet(path: Path):
    doc = Document()
    configure_document(doc, "Springer Journal Author Details Worksheet")
    add_title_block(
        doc,
        "Conference author support pack",
        "Author Details & Declarations Worksheet",
        "For papers considered for the selected Springer journals after GLOGIFT 2027",
        "Editable working file  |  Source check: 1 August 2026",
    )
    add_callout(
        doc,
        "Important",
        "This is a GLOGIFT 2027 preparation aid, not a publisher-issued Springer template. Keep this file separate from the blinded manuscript. Enter the same information in the journal submission system wherever requested.",
    )

    doc.add_heading("1. Target journal", level=1)
    add_field_table(
        doc,
        [
            ("Selected journal", "[Choose: Global Journal of Flexible Systems Management / International Journal of Global Business and Competitiveness]"),
            ("Article type", "[Original research / conceptual paper / review / case or other accepted type]"),
            ("GLOGIFT paper ID", "[Paper ID]"),
        ],
    )

    doc.add_heading("2. Manuscript title", level=1)
    add_field_table(doc, [("Full title", "[Concise and informative article title]"), ("Short title", "[Optional running title]")])

    doc.add_heading("3. Author information", level=1)
    p = doc.add_paragraph()
    set_run_font(p.add_run("Repeat the block below for every author in the final publication order."), italic=True, color=MUTED)
    add_field_table(
        doc,
        [
            ("Author number", "[1 / 2 / 3 / ...]"),
            ("Full name", "[Given name, middle name/initial, family name]"),
            ("Affiliation", "[Department, institution, city, state, country]"),
            ("E-mail", "[Active institutional or professional e-mail]"),
            ("ORCID", "[https://orcid.org/0000-0000-0000-0000, if available]"),
            ("Corresponding author", "[Yes / No]"),
        ],
    )

    doc.add_heading("4. Corresponding author", level=1)
    add_field_table(
        doc,
        [
            ("Name", "[Full name]"),
            ("E-mail", "[Active e-mail]"),
            ("Postal address", "[Complete correspondence address, if requested by the journal]"),
        ],
    )

    doc.add_heading("5. Author contributions", level=1)
    p = doc.add_paragraph()
    set_run_font(p.add_run("[List each author’s specific contribution. A CRediT-style statement may be used: conceptualization; methodology; investigation; formal analysis; data curation; writing—original draft; writing—review and editing; supervision; funding acquisition.]"), color=INK)

    doc.add_heading("6. Acknowledgements, funding and interests", level=1)
    add_field_table(
        doc,
        [
            ("Acknowledgements", "[People, institutional support and non-author contributions; spell out organization names]"),
            ("Funding", "[Funder name and grant number, or: The authors received no funding for this work.]"),
            ("Competing interests", "[Describe financial/non-financial interests, or: The authors have no relevant financial or non-financial interests to disclose.]"),
        ],
    )

    doc.add_heading("7. Ethics, consent, data and AI use", level=1)
    add_field_table(
        doc,
        [
            ("Ethics approval", "[Approval body and reference number / Not applicable]"),
            ("Consent", "[Informed consent statement / Not applicable]"),
            ("Data availability", "[Repository and persistent link, access conditions, or justified restriction]"),
            ("AI-tool disclosure", "[Describe generative AI use in the Methods or suitable section. AI-assisted copy editing alone need not be declared under the current journal guidance.]"),
        ],
    )

    doc.add_heading("8. Final author confirmation", level=1)
    for statement in (
        "All authors approve the manuscript, author order and corresponding author designation.",
        "The work has not been published previously and is not under consideration elsewhere.",
        "Necessary permissions have been obtained for reused text, tables, figures or other material.",
        "Every author accepts accountability for the final manuscript; no large language model is listed as an author.",
    ):
        add_bullet(doc, statement)

    add_source_list(doc)
    doc.save(path)


def add_placeholder_section(doc: Document, heading: str, guidance: str, sample_lines=2):
    doc.add_heading(heading, level=1 if heading.split(".")[0].isdigit() else 2)
    p = doc.add_paragraph()
    set_run_font(p.add_run(guidance), size=10, color=MUTED, italic=True)
    for _ in range(sample_lines):
        p = doc.add_paragraph()
        set_run_font(p.add_run("[Replace this instructional text with manuscript content.]"), color="6B7280")


def create_blinded_manuscript_template(path: Path):
    doc = Document()
    configure_document(doc, "Springer Journal Blinded Manuscript Template")
    add_title_block(
        doc,
        "Conference author support pack",
        "Blinded Manuscript Template",
        "For the two selected Springer journals associated with GLOGIFT 2027",
        "Editable Word template  |  Double-anonymous review",
    )
    add_callout(
        doc,
        "Before writing",
        "Delete all instructional text in brackets. Do not include author names, affiliations, acknowledgements or identity-revealing file metadata in this blinded manuscript. Use the separate Author Details & Declarations Worksheet for those items.",
        fill=LIGHT_GOLD,
    )
    add_callout(
        doc,
        "Journal-specific rule",
        "For the International Journal of Global Business and Competitiveness (IJGBC), keep the complete manuscript below 4,000 words including tables, figures and references, and retain Section 7 with 4–5 application-oriented questions. For the Global Journal of Flexible Systems Management (GJFSM), delete Section 7 unless the editor or call requires it.",
        fill=LIGHT_BLUE,
        accent=BLUE,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(8)
    set_run_font(p.add_run("[Concise and Informative Article Title]"), size=18, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("Target journal: [GJFSM / IJGBC]  |  Manuscript word count: [0000]"), size=10, color=MUTED)

    doc.add_heading("Abstract", level=1)
    p = doc.add_paragraph()
    set_run_font(p.add_run("[Write 150–250 words. State the purpose, approach/method, principal findings and contribution or implications. Do not use undefined abbreviations or unspecified references.]"), color="6B7280")

    doc.add_heading("Keywords", level=2)
    p = doc.add_paragraph()
    set_run_font(p.add_run("[Provide 4–6 indexing terms, separated by semicolons.]"), color="6B7280")

    doc.add_heading("JEL classification codes", level=2)
    p = doc.add_paragraph()
    set_run_font(p.add_run("[Insert appropriate JEL code(s), e.g., M15; O33. Verify against the current classification.]"), color="6B7280")

    add_placeholder_section(doc, "1. Introduction", "Establish the managerial problem, research gap, objective and contribution. Explain journal fit without identifying the authors.", 3)
    add_placeholder_section(doc, "2. Literature Review and Conceptual Background", "Synthesize relevant theory and prior findings. State propositions or hypotheses where applicable. Use author–date citations.", 3)
    add_placeholder_section(doc, "3. Methodology", "Explain the research design, setting, sample/data, measures, procedures and analysis sufficiently for evaluation and replication. Document substantive generative-AI use here or in a suitable alternative section.", 3)

    doc.add_heading("3.1 Data and sample", level=2)
    p = doc.add_paragraph()
    set_run_font(p.add_run("[Describe data sources, sampling, inclusion criteria and final analytical sample.]"), color="6B7280")
    doc.add_heading("3.2 Measures and analysis", level=2)
    p = doc.add_paragraph()
    set_run_font(p.add_run("[Define constructs, variables, analytical procedures, validity and robustness checks.]"), color="6B7280")

    add_placeholder_section(doc, "4. Results / Findings", "Report findings in a logical order. Cite every table and figure consecutively in the text and avoid duplicating the same data in prose and tables.", 3)

    doc.add_heading("Table 1  [Concise table title]", level=2)
    table = doc.add_table(rows=4, cols=3)
    set_table_geometry(table, [2400, 3480, 3480])
    set_table_borders(table)
    mark_table_header(table.rows[0])
    for j, heading in enumerate(("Construct / category", "Measure / evidence", "Result / interpretation")):
        set_cell_shading(table.rows[0].cells[j], LIGHT_BLUE)
        p = table.rows[0].cells[j].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(heading), size=9.5, color=NAVY, bold=True)
    for row in table.rows[1:]:
        for cell in row.cells:
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run("[Insert data]"), size=9.5, color=MUTED)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    set_run_font(p.add_run("Note. [Define abbreviations, significance markers and source/permission information where applicable.]"), size=9, color=MUTED, italic=True)

    doc.add_heading("Fig. 1  [Concise figure caption]", level=2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    set_run_font(p.add_run("[Insert figure here. Ensure readable lettering, accessible contrast and a descriptive caption.]"), size=10, color=MUTED, italic=True)

    add_placeholder_section(doc, "5. Discussion", "Interpret the findings in relation to the research questions and prior literature. Explain theoretical and practical significance, boundary conditions and unexpected results.", 3)
    add_placeholder_section(doc, "6. Conclusions, Implications and Limitations", "State the main conclusion, managerial implications, limitations and focused directions for future research.", 3)
    add_placeholder_section(doc, "7. Key Questions Reflecting Applicability in Real Life (IJGBC only)", "For IJGBC, provide 4–5 forward-looking questions or tasks that encourage practitioners/researchers to apply the paper’s learning. Avoid closed yes/no questions. For GJFSM, delete this section unless requested.", 1)
    for i in range(1, 6):
        add_number(doc, f"[Application-oriented question {i}]" )

    doc.add_heading("Statements and Declarations", level=1)
    for subheading, guidance in (
        ("Funding", "[State the funder and grant number, or state that no funding was received.]"),
        ("Competing interests", "[Disclose relevant financial/non-financial interests, or state that none exist.]"),
        ("Ethics approval and consent", "[Provide approval/consent information where applicable, without revealing author identity during review; otherwise state Not applicable.]"),
        ("Data availability", "[For original research, explain how supporting data can be accessed, including repository links/identifiers or justified access restrictions. Use an anonymized link during review where necessary.]"),
        ("Generative AI use", "[If applicable, document substantive use of an LLM or other generative AI tool. Do not list an AI tool as an author.]"),
    ):
        doc.add_heading(subheading, level=2)
        p = doc.add_paragraph()
        set_run_font(p.add_run(guidance), color="6B7280")

    doc.add_heading("References", level=1)
    p = doc.add_paragraph()
    set_run_font(p.add_run("[List only cited works. Arrange alphabetically by first-author surname. Use consistent APA 7 author–date style; italicize journal names and book titles; include DOI links when available.]"), color="6B7280")
    p = doc.add_paragraph()
    set_run_font(p.add_run("Example journal article: Author, A. A., Author, B. B., & Author, C. C. (Year). Article title. "), color=INK)
    set_run_font(p.add_run("Journal Title"), color=INK, italic=True)
    set_run_font(p.add_run(", volume(issue), pages. https://doi.org/xxxxx"), color=INK)

    doc.add_page_break()
    doc.add_heading("Final blinded-file check", level=1)
    for statement in (
        "No author names, affiliations, acknowledgements, self-identifying notes or identity-revealing document metadata remain.",
        "Abstract is 150–250 words; 4–6 keywords and appropriate JEL codes are present.",
        "Headings use decimal numbering and no more than three levels.",
        "References use author–date citations and an alphabetized APA-style list with DOI links where available.",
        "Tables and figures are numbered, cited consecutively, captioned and supplied in editable/high-quality form.",
        "Required statements and a data availability statement (for original research) are included or completed in the submission interface.",
        "IJGBC only: total manuscript is under 4,000 words and Section 7 contains 4–5 application-oriented questions.",
        "All bracketed instructions have been deleted before submission.",
    ):
        add_bullet(doc, statement)
    add_source_list(doc)
    doc.save(path)


def create_guidelines(path: Path):
    doc = Document()
    configure_document(doc, "Springer Journal Submission Guidelines")
    add_title_block(
        doc,
        "Conference author support pack",
        "Paper Submission Guidelines",
        "For GLOGIFT 2027 papers considered for the two selected Springer journals",
        "Prepared from official journal pages only  |  Checked 1 August 2026",
    )
    add_callout(
        doc,
        "Scope and authority",
        "This guide summarizes the current official journal instructions for conference authors. It is not an acceptance promise, a publisher-issued template or a substitute for the live journal page. Selected papers remain subject to independent editorial screening, peer review and revision.",
    )

    doc.add_heading("1. Select the correct target journal", level=1)
    table = doc.add_table(rows=3, cols=3)
    set_table_geometry(table, [2200, 3580, 3580])
    set_table_borders(table)
    mark_table_header(table.rows[0])
    headers = ("Journal", "Best fit", "Journal-specific requirement")
    for i, value in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], NAVY)
        p = table.rows[0].cells[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(value), size=9.5, color=WHITE, bold=True)
    rows = (
        (
            "Global Journal of Flexible Systems Management (GJFSM)",
            "General management research centred on organizational flexibility—adaptive, responsive and agile strategy, structure, systems, people and culture; includes conceptual, empirical, case, methodological and review work.",
            "No special manuscript word limit is stated on the current journal page. Follow the common Springer requirements and re-check the live page before submission.",
        ),
        (
            "International Journal of Global Business and Competitiveness (IJGBC)",
            "Research on management and business practices that strengthen international competitiveness, including sustainability, flexibility, digital innovation, internationalization and strategic leadership.",
            "Fewer than 4,000 words including tables, figures and references; add 4–5 'Key Questions Reflecting Applicability in Real Life' after Conclusions and before References.",
        ),
    )
    for r_index, row_data in enumerate(rows, start=1):
        for c_index, value in enumerate(row_data):
            cell = table.rows[r_index].cells[c_index]
            if r_index % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run(value), size=9.3, color=INK, bold=(c_index == 0))

    doc.add_heading("2. Required submission files", level=1)
    for item in (
        "Blinded manuscript file (.docx or LaTeX source; include editable source files at every submission and revision).",
        "Separate author-details/title-page information or the equivalent fields in the journal submission system, as instructed during submission.",
        "Editable tables and figures, plus separate high-resolution artwork files if required by size or upload constraints.",
        "Supplementary files in standard formats, clearly named and described; for IJGBC, overflow appendix material should be uploaded as a separate file named 'Appendix'.",
        "Permissions evidence for any reused text, table, figure or other copyrighted material.",
    ):
        add_bullet(doc, item)

    doc.add_heading("3. Double-anonymous peer review", level=1)
    for item in (
        "Remove author names, affiliations and all potentially identifying information from the manuscript and accompanying files.",
        "Keep title-page details, acknowledgements, funding and identity-bearing disclosures in the separate author worksheet/file or submission-system fields.",
        "Avoid self-citations phrased in a way that reveals identity; cite prior work neutrally in the third person.",
        "Anonymize repository links, supplementary files and file metadata where they could identify the authors.",
    ):
        add_bullet(doc, item)

    doc.add_heading("4. Manuscript structure and formatting", level=1)
    structure = (
        ("Title", "Concise and informative; no author information in the blinded file."),
        ("Abstract", "150–250 words; no undefined abbreviations or unspecified references."),
        ("Keywords", "4–6 indexing terms."),
        ("JEL codes", "Provide an appropriate number of Journal of Economic Literature classification codes."),
        ("Main text", "Use decimal headings with no more than three levels; define abbreviations at first use."),
        ("Footnotes", "Use footnotes, not endnotes; number text footnotes consecutively; do not use them as a substitute for the reference list."),
        ("Acknowledgements", "Keep on the separate title page/author-information file; write funding organization names in full."),
        ("Source format", "Springer recommends its LaTeX template; Word files are also accepted."),
    )
    add_field_table(doc, structure, widths=(2100, 7260))

    doc.add_heading("5. IJGBC-specific requirements", level=1)
    add_callout(doc, "Manuscript length", "The complete manuscript must be fewer than 4,000 words, including tables, figures and references. Move less-important material to a separately uploaded supplementary Appendix when necessary.", fill=LIGHT_BLUE, accent=BLUE)
    add_callout(doc, "Application questions", "After Conclusions and before References, add 4–5 forward-looking questions or tasks under 'Key Questions Reflecting Applicability in Real Life'. They should prompt application by practitioners or researchers rather than invite a closed yes/no response.", fill=LIGHT_BLUE, accent=BLUE)
    add_bullet(doc, "List the specific contribution of each author in the author-details file and required system field.")

    doc.add_heading("6. References", level=1)
    for item in (
        "Use author–date citations in the text, for example: (Author, 2025) or Author (2025).",
        "Follow APA 7 guidance consistently; include all authors up to 20 where practical, in line with the journal page.",
        "Alphabetize the reference list by the surname of the first author.",
        "Include only works cited in the manuscript and published or accepted for publication.",
        "Italicize journal names and book titles; include DOI links in full URL form whenever available.",
    ):
        add_bullet(doc, item)

    doc.add_heading("7. Tables and figures", level=1)
    table = doc.add_table(rows=6, cols=2)
    set_table_geometry(table, [2500, 6860])
    set_table_borders(table)
    items = (
        ("Tables", "Number with Arabic numerals, cite consecutively in the text, provide a concise caption and identify reused sources in the caption."),
        ("Figure files", "Preferred: EPS for vectors and TIFF for halftones; MS Office files are accepted. Embed fonts in vector graphics and name files Fig1, Fig2, etc."),
        ("Resolution", "Minimum 1200 dpi for bitmap line art, 300 dpi for halftones and 600 dpi for combination artwork."),
        ("Lettering", "Use readable Helvetica/Arial-style lettering, generally 8–12 pt at final size; avoid decorative effects and embedded titles/captions."),
        ("Colour/accessibility", "Use RGB; ensure information remains distinguishable without colour, use patterns where helpful and provide descriptive captions with adequate contrast."),
        ("Placement", "Place figures within the manuscript body unless file size requires separate upload; cite every figure in consecutive order."),
    )
    for r_index, (label, detail) in enumerate(items):
        left, right = table.rows[r_index].cells
        set_cell_shading(left, LIGHT_BLUE)
        lp = left.paragraphs[0]
        rp = right.paragraphs[0]
        lp.paragraph_format.space_after = Pt(0)
        rp.paragraph_format.space_after = Pt(0)
        set_run_font(lp.add_run(label), size=9.5, color=NAVY, bold=True)
        set_run_font(rp.add_run(detail), size=9.5, color=INK)

    doc.add_heading("8. Statements, ethics, data and AI use", level=1)
    for item in (
        "Competing interests: disclose relevant financial and non-financial interests; submissions lacking relevant declarations may be returned as incomplete.",
        "Funding and acknowledgements: identify funders and grant numbers, or state that no funding was received.",
        "Author contributions: specify each author’s contribution; the corresponding author confirms author approval, order and accountability.",
        "Ethics and consent: include approvals and informed-consent statements when relevant to the research.",
        "Data availability: every original research article must explain how supporting data can be accessed, or describe justified restrictions and reuse conditions.",
        "Generative AI: large language models cannot be authors. Substantive LLM/AI use must be documented in Methods or a suitable section; AI-assisted copy editing alone need not be declared under the current journal guidance. Human authors remain accountable.",
    ):
        add_bullet(doc, item)

    doc.add_heading("9. Pre-submission checklist", level=1)
    for item in (
        "The paper fits the aims and scope of the selected journal.",
        "All authors approve the manuscript, author order and journal submission.",
        "The manuscript is original and is not being considered by another journal.",
        "The blinded file contains no identifying information or revealing metadata.",
        "Abstract, keywords, JEL codes, numbered headings, references and declarations are complete.",
        "All tables and figures are cited, captioned, legible and supplied in suitable editable/high-resolution formats.",
        "Permissions are attached for reused material.",
        "Original research includes a data availability statement.",
        "IJGBC only: the manuscript is under 4,000 words and includes 4–5 application questions in the required position.",
        "The live official submission page has been re-checked immediately before upload.",
    ):
        add_bullet(doc, item)

    doc.add_heading("10. Suggested file names", level=1)
    add_field_table(
        doc,
        [
            ("Blinded manuscript", "GLOGIFT2027_PaperID_Blinded_Manuscript.docx"),
            ("Author details", "GLOGIFT2027_PaperID_Author_Details.docx"),
            ("Figures", "Fig1.eps, Fig2.tif, ..."),
            ("IJGBC appendix", "Appendix.pdf or the standard file format required by the content"),
        ],
    )

    add_source_list(doc)
    doc.save(path)


def audit_docx(path: Path):
    doc = Document(path)
    assert len(doc.sections) >= 1
    section = doc.sections[0]
    assert round(section.page_width.inches, 2) == 8.50
    assert round(section.page_height.inches, 2) == 11.00
    assert round(section.left_margin.inches, 2) == 1.00
    assert round(section.right_margin.inches, 2) == 1.00
    for table in doc.tables:
        grid = table._tbl.tblGrid
        widths = [int(col.get(qn("w:w"))) for col in grid.findall(qn("w:gridCol"))]
        assert widths and sum(widths) == 9360, (path.name, widths)
        tbl_w = table._tbl.tblPr.first_child_found_in("w:tblW")
        assert int(tbl_w.get(qn("w:w"))) == 9360
        assert tbl_w.get(qn("w:type")) == "dxa"
        tbl_ind = table._tbl.tblPr.first_child_found_in("w:tblInd")
        assert int(tbl_ind.get(qn("w:w"))) == 120


def main():
    OUTPUT.mkdir(parents=True, exist_ok=True)
    files = [
        OUTPUT / "01_GLOGIFT2027_Springer_Author_Details_Worksheet.docx",
        OUTPUT / "02_GLOGIFT2027_Springer_Blinded_Manuscript_Template.docx",
        OUTPUT / "03_GLOGIFT2027_Springer_Submission_Guidelines.docx",
    ]
    create_author_details_worksheet(files[0])
    create_blinded_manuscript_template(files[1])
    create_guidelines(files[2])
    for path in files:
        audit_docx(path)
        print(path)


if __name__ == "__main__":
    main()
